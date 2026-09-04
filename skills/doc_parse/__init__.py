"""doc_parse：解析 Word 文档并提取文本"""
from skills import Skill


class DocParse(Skill):
    name = "doc_parse"
    description = "解析 Word 文档并提取文本内容（段落 + 表格）"
    input_schema = {"file_path": {"type": "string", "description": "Word 文档路径"}, "max_chars": {"type": "integer", "description": "最大返回字符数"}}

    def execute(self, file_path="", max_chars=3000, **kwargs):
        if not file_path:
            return {"results": [], "error": "缺少 file_path 参数"}
        try:
            from docx import Document
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            text = "\n".join(parts)
            return {"results": [{"content": text[:max_chars]}], "total": len(parts)}
        except Exception as e:
            return {"results": [], "error": str(e)[:200]}
